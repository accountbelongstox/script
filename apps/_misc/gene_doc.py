# from win32com import client as wc
from pycore.requirement_fn.auto_install import auto_install
import subprocess

if __name__ == "__main__":
    auto_install.start()

    import os
    from docx import Document
    from pycore.utils import file
    import time
    from win32com.client import Dispatch
    import win32com.client

    class WordProcessor:
        max_size = 480 * 1024 * 1024  # 500 MB
        current_insert_doc = None
        document_count = 0

        def __init__(self):
            self.output_folder = os.path.join(file.get_root_dir(), "out/doc")
            os.makedirs(self.output_folder, exist_ok=True)

        def generate_new_document(self):
            document_name = str(int(time.time()))
            new_document_path = os.path.join(self.output_folder, f"{document_name}.docx")
            document = Document()
            document.save(new_document_path)
            print(f"New document created: {self.basename(new_document_path)}")
            return new_document_path

        def get_file_size(self, file_path):
            try:
                size = os.path.getsize(file_path)
                print(f"File size of {file_path}: {size} bytes")
                return size
            except FileNotFoundError:
                print(f"File not found: {file_path}")
                return 0

        def get_current_doc(self):
            if self.current_insert_doc is None:
                self.current_insert_doc = self.generate_new_document()
            if self.get_file_size(self.current_insert_doc) > self.max_size:
                self.current_insert_doc = self.generate_new_document()
            try:
                return Document(self.current_insert_doc)
            except Exception as e:
                print(e)
            return

        def read_word_document(self, file_path):
            if file_path.lower().endswith('.docx'):
                return self.read_docx_document(file_path)
            elif file_path.lower().endswith('.doc'):
                return self.read_doc_document(file_path)
            else:
                print(f"Unsupported file format: {self.basename(file_path)}")
                return None

        def basename(self, path):
            return os.path.basename(path)

        def read_docx_document(self, file_path):
            try:
                doc = Document(file_path)
                content = [paragraph.text for paragraph in doc.paragraphs]
                print(f"Read content from {self.basename(file_path)}: {len(content)}")
                return content
            except Exception as e:
                print(e)
                return None

        def read_doc_document(self, file_path):
            try:
                wordapp = win32com.client.gencache.EnsureDispatch("Word.Application")
                wordapp.Visible = 0
                worddoc = wordapp.Documents.Open(file_path)
                wdFormatText = 2
                basename = os.path.basename(file_path)
                basename = os.path.splitext(basename)[0] + ".txt"
                txt_file_path = file.resolve_path(basename,"out/tmp/doc")
                file.mkbasedir(txt_file_path)
                worddoc.SaveAs(txt_file_path, FileFormat=wdFormatText)
                worddoc.Close()
                wordapp.Quit()
                txt_content = file.read_text(txt_file_path,encoding="gbk")
                return txt_content
            except Exception as e:
                print("Exception:", e)
                return None
            # try:
            #     doc = word_app.Documents.Open(file_path)
            #     content = [paragraph.Text for paragraph in doc.Paragraphs]
            #     print(f"Read content from {self.basename(file_path)}: {len(content)}")
            #     doc.Close()
            #     word_app.Documents.Close(wc.wdDoNotSaveChanges)
            #     word_app.Quit()
            #     return content
            # except Exception as e:
            #     print(f"Error reading .doc file: {e}")
            #     return None

        def combine_documents(self, source_folder):
            for filename in os.listdir(source_folder):
                if filename.lower().endswith('.docx'):
                    current_document = self.get_current_doc()
                    file_path = os.path.join(source_folder, filename)
                    content = self.read_docx_document(file_path)
                    if content:
                        current_document.add_paragraph("\n".join(content))
                        current_document.save(self.current_insert_doc)
                elif filename.lower().endswith('.doc'):
                    current_document = self.get_current_doc()
                    file_path = os.path.join(source_folder, filename)
                    content = self.read_doc_document(file_path)
                    if content:
                        current_document.add_paragraph("\n".join(content))
                        current_document.save(self.current_insert_doc)
            print(f"Combined {self.document_count} documents.")


    word_processor = WordProcessor()
    source_folder = "out/testfiles"
    word_processor.combine_documents(source_folder)
